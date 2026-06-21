from docx import Document
import os

# 读取文本文件内容
txt_path = "c:\\Users\\Linxiaorou\\Desktop\\website\\qq音乐\\课程设计报告内容.txt"
docx_path = "c:\\Users\\Linxiaorou\\Desktop\\website\\qq音乐\\课程设计报告-Web应用开发.docx"

# 创建一个新的Word文档
doc = Document()

# 设置文档标题
doc.add_heading('QQ音乐Web应用开发课程设计报告', level=1)

# 读取并解析文本内容
with open(txt_path, 'r', encoding='utf-8') as f:
    content = f.readlines()

current_heading_level = 1
current_paragraph = []

for line in content:
    line = line.strip()
    if not line:
        # 处理空行，添加当前段落并重置
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
            current_paragraph = []
        continue
    
    # 处理标题
    if line.startswith('#'):
        # 添加当前段落（如果有）
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
            current_paragraph = []
        
        # 确定标题级别
        heading_level = len(line.split()[0])
        heading_text = line[heading_level+1:].strip()
        doc.add_heading(heading_text, level=heading_level)
        current_heading_level = heading_level
    
    # 处理列表项
    elif line.startswith('- '):
        # 添加当前段落（如果有）
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
            current_paragraph = []
        
        # 添加列表项
        doc.add_paragraph(line[2:], style='List Bullet')
    
    # 处理表格
    elif line.startswith('|'):
        # 添加当前段落（如果有）
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
            current_paragraph = []
        
        # 解析表格行
        row_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        
        # 检查是否是表头
        if any('---' in cell for cell in row_cells):
            # 跳过表头分隔线
            continue
        
        # 添加表格（如果是第一行）或表格行
        if not hasattr(doc, '_table_started'):
            doc._table_started = True
            doc._current_table = doc.add_table(rows=1, cols=len(row_cells))
            for i, cell_text in enumerate(row_cells):
                doc._current_table.cell(0, i).text = cell_text
        else:
            row = doc._current_table.add_row()
            for i, cell_text in enumerate(row_cells):
                row.cells[i].text = cell_text
    
    # 处理代码块
    elif line.startswith('```'):
        # 添加当前段落（如果有）
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
            current_paragraph = []
        
        # 检查是否是代码块结束
        if hasattr(doc, '_code_block_started'):
            doc._code_block_started = False
            # 添加代码块结束
            doc.add_paragraph('```')
        else:
            doc._code_block_started = True
            # 添加代码块开始
            doc.add_paragraph('```')
    
    # 处理普通文本
    else:
        if hasattr(doc, '_code_block_started') and doc._code_block_started:
            # 如果在代码块中，直接添加段落
            doc.add_paragraph(line)
        else:
            # 否则添加到当前段落
            current_paragraph.append(line)

# 添加最后一个段落（如果有）
if current_paragraph:
    doc.add_paragraph(' '.join(current_paragraph))

# 保存文档
doc.save(docx_path)
print(f"报告内容已成功写入 {docx_path}")
